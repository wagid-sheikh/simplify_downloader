"""
Builds a resumable, grouped, editable Word reference document from CharmsWiki
pages, ordered for system design and SRS authoring.
"""

import asyncio
import csv
import json
import random
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------- Paths (all relative to this script file) ----------
BASE_DIR = Path(__file__).resolve().parent

RAW_CSV = BASE_DIR / "wiki_links.csv"             # exists already (not regenerated)
GROUPED_CSV = BASE_DIR / "wiki_links_grouped.csv" # exists already (design order)

OUT_DIR = BASE_DIR / "Wiki-PDF"  # keep same folder name you already use
CACHE_DIR = OUT_DIR / "_cache_pages_text"         # per-page cached text
STATE_PATH = OUT_DIR / "merge_state.json"         # resume marker (kept)
DOCX_OUT = OUT_DIR / "CharmsWiki_REFERENCE.docx"
DOCX_TMP = OUT_DIR / "CharmsWiki_REFERENCE.tmp.docx"


# ---------- Controls ----------
NAV_TIMEOUT_MS = 45_000
RETRY_COUNT = 2

# Validation run: keep 2 for now, then set to None for full run (388+)
MAX_TO_SAVE = 2

# Throttle delay between pages
DELAY_MIN = 15
DELAY_MAX = 20


@dataclass
class LinkRow:
    design_seq: int
    seq: int
    group: str
    subgroup: str
    name: str
    url: str


def chrome_executable_path() -> str:
    """
    Prefer explicit override, else use standard macOS Chrome path.
    Set CHROME_PATH if needed.
    """
    import os
    env = os.environ.get("CHROME_PATH")
    if env and Path(env).exists():
        return env

    candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Google Chrome Canary.app/Contents/MacOS/Google Chrome Canary",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
    ]
    for p in candidates:
        if Path(p).exists():
            return p

    raise FileNotFoundError(
        "Could not find a local Chrome/Canary/Chromium executable. "
        "Install Google Chrome, or set CHROME_PATH to the full executable path."
    )


def sanitize_filename(name: str, max_len: int = 140) -> str:
    name = (name or "").strip()
    name = re.sub(r"\s+", " ", name)
    name = re.sub(r'[\/\\\:\*\?\"\<\>\|\n\r\t]', "-", name)
    name = re.sub(r"[\u0000-\u001f]", "", name)
    name = name.strip(" .-_")
    if not name:
        name = "untitled"
    if len(name) > max_len:
        name = name[:max_len].rstrip(" .-_")
    return name


def cache_text_path(design_seq: int, label: str) -> Path:
    safe = sanitize_filename(label, max_len=120)
    return CACHE_DIR / f"{design_seq:04d} - {safe}.json"


def load_state() -> dict:
    if STATE_PATH.exists():
        return json.loads(STATE_PATH.read_text(encoding="utf-8"))
    return {"last_completed_design_seq": 0, "cached_design_seqs": []}


def save_state(state: dict) -> None:
    STATE_PATH.write_text(json.dumps(state, indent=2), encoding="utf-8")


def read_grouped_csv(path: Path) -> list[LinkRow]:
    """
    Expected header:
    Design Sequence,Sequence,Group,Subgroup,Link Name,URL
    """
    if not path.exists():
        raise FileNotFoundError(f"Missing required file: {path}")

    rows: list[LinkRow] = []
    with path.open("r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)

        required = {"Design Sequence", "Sequence", "Group", "Subgroup", "Link Name", "URL"}
        got = set(r.fieldnames or [])
        if required != got:
            raise ValueError(
                f"Grouped CSV header mismatch.\nExpected: {required}\nGot:      {got}\nFile: {path}"
            )

        for line in r:
            rows.append(
                LinkRow(
                    design_seq=int(line["Design Sequence"]),
                    seq=int(line["Sequence"]),
                    group=(line["Group"] or "").strip(),
                    subgroup=(line["Subgroup"] or "").strip(),
                    name=(line["Link Name"] or "").strip(),
                    url=(line["URL"] or "").strip(),
                )
            )

    rows.sort(key=lambda x: x.design_seq)
    return [x for x in rows if x.name and x.url]


async def fetch_page_text(context, row: LinkRow, idx: int, total: int) -> dict:
    """
    Fetch page and return structured cache payload.
    Idempotency: caller decides whether to fetch based on cached file existence.
    """
    for attempt in range(RETRY_COUNT + 1):
        page = await context.new_page()
        try:
            print(f"[{idx}/{total}] FETCH  {row.group} > {row.subgroup} :: {row.name}")
            await page.goto(row.url, wait_until="networkidle", timeout=NAV_TIMEOUT_MS)

            # Try to capture the core readable content.
            # We keep it simple and robust: use body innerText.
            body_text = await page.evaluate("() => document.body ? document.body.innerText : ''")
            page_title = await page.title()

            payload = {
                "design_seq": row.design_seq,
                "seq": row.seq,
                "group": row.group,
                "subgroup": row.subgroup,
                "name": row.name,
                "url": row.url,
                "page_title": page_title or "",
                "text": body_text or "",
            }

            await page.close()
            return payload

        except PlaywrightTimeoutError:
            await page.close()
            if attempt >= RETRY_COUNT:
                raise
            print(f"[{idx}/{total}] Timeout, retrying ({attempt+1}/{RETRY_COUNT})...")
            await asyncio.sleep(1.5)

        except Exception:
            await page.close()
            if attempt >= RETRY_COUNT:
                raise
            print(f"[{idx}/{total}] Error, retrying ({attempt+1}/{RETRY_COUNT})...")
            await asyncio.sleep(1.5)


def write_cache(payload: dict) -> Path:
    p = cache_text_path(payload["design_seq"], payload["name"])
    p.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return p


def load_all_cached_payloads(rows: list[LinkRow]) -> list[dict]:
    """
    Load cached payloads for rows that exist in cache.
    Returns in design_seq order (the same as grouped CSV).
    """
    cached = []
    for row in rows:
        # Find file by deterministic name pattern; we compute expected path.
        p = cache_text_path(row.design_seq, row.name)
        if p.exists():
            cached.append(json.loads(p.read_text(encoding="utf-8")))
    cached.sort(key=lambda x: int(x.get("design_seq", 0)))
    return cached


def add_toc_field(paragraph):
    """
    Insert a TOC field that Word can update.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Table of Contents (Right-click → Update Field)"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def build_docx_from_cache(rows: list[LinkRow], cached_payloads: list[dict], out_path: Path) -> None:
    """
    Build the main DOCX from cached data (idempotent build).
    TOC will be inserted afterward (per your requirement).
    """
    doc = Document()

    # Title
    doc.add_heading("CharmsWiki Reference Binder (Grouped)", level=0)

    # Intro paragraph (as requested)
    intro = (
        "This document is an automatically compiled reference binder created by visiting the CharmsWiki pages "
        "in a system-design-oriented order (grouped by modules such as Platform & Access, Case Management, "
        "Carer Lifecycle, Placements, Finance, Reporting, and Documents). "
        "Use it alongside the Blueprint to validate requirements, terminology, workflows, and report definitions. "
        "The Table of Contents is generated from headings—open this file in Microsoft Word and Update the TOC "
        "to refresh page numbers after any edits."
    )
    doc.add_paragraph(intro)

    doc.add_paragraph("Source inventory (authoritative order): wiki_links_grouped.csv")
    doc.add_paragraph(f"Total pages intended: {len(rows)}")
    doc.add_page_break()

    # Content in grouped order:
    # Heading 1: Group
    # Heading 2: Subgroup
    # Heading 3: Page name
    current_group = None
    current_subgroup = None

    payload_by_seq = {int(p["design_seq"]): p for p in cached_payloads}

    for row in rows:
        p = payload_by_seq.get(row.design_seq)
        if not p:
            # Not cached yet; skip content but keep structure optional
            # For now, we will include a placeholder entry
            if row.group != current_group:
                current_group = row.group
                doc.add_heading(current_group, level=1)
                current_subgroup = None

            if row.subgroup != current_subgroup:
                current_subgroup = row.subgroup
                doc.add_heading(current_subgroup, level=2)

            doc.add_heading(f"{row.design_seq:04d}. {row.name}", level=3)
            doc.add_paragraph(f"URL: {row.url}")
            doc.add_paragraph("Status: Not cached yet (run script again to fetch).")
            doc.add_page_break()
            continue

        if p.get("group") != current_group:
            current_group = p.get("group") or "Unclassified"
            doc.add_heading(current_group, level=1)
            current_subgroup = None

        if p.get("subgroup") != current_subgroup:
            current_subgroup = p.get("subgroup") or "Other"
            doc.add_heading(current_subgroup, level=2)

        doc.add_heading(f"{int(p['design_seq']):04d}. {p.get('name','')}", level=3)
        doc.add_paragraph(f"URL: {p.get('url','')}")
        title = (p.get("page_title") or "").strip()
        if title:
            doc.add_paragraph(f"Page Title: {title}")

        text = (p.get("text") or "").strip()
        if not text:
            doc.add_paragraph("(No text extracted.)")
        else:
            # Split into paragraphs for Word readability.
            for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
                doc.add_paragraph(block)

        doc.add_page_break()

    doc.save(str(out_path))


def insert_toc_at_front(docx_path: Path) -> None:
    """
    Your requirement:
    - produce TOC toward the end (we do it after building content),
    - then insert it back into the final document.
    Here: we open the doc and insert a TOC section near the start.
    """
    doc = Document(str(docx_path))

    # Insert TOC after the second paragraph (after title + intro) for best usability.
    # python-docx has limited true insertion; we do a pragmatic method:
    # create TOC at end of intro area by inserting paragraphs near start using XML.
    # We'll insert after paragraph index 2 if possible.

    insert_at = 2 if len(doc.paragraphs) >= 2 else len(doc.paragraphs)

    # Create a new paragraph for TOC heading
    toc_heading = doc.add_paragraph()
    toc_heading.style = doc.styles["Heading 1"]
    toc_heading.add_run("Table of Contents")

    toc_field_para = doc.add_paragraph()
    add_toc_field(toc_field_para)

    doc.add_paragraph("Note: In Word, right-click the Table of Contents → Update Field → Update entire table.")
    doc.add_page_break()

    # Move these last inserted paragraphs to the front region by XML manipulation.
    # We appended them at the end; now we relocate them.
    body = doc._body._element
    # Grab the last 4 block elements we just added: heading, field para, note para, page break para
    moved = [body[-4], body[-3], body[-2], body[-1]]

    # Insert them after the chosen paragraph position in the body
    # Find the XML element corresponding to the insert_at paragraph
    if insert_at < len(doc.paragraphs):
        anchor_p = doc.paragraphs[insert_at]._p
        anchor_idx = list(body).index(anchor_p)
        for i, el in enumerate(moved):
            body.insert(anchor_idx + 1 + i, el)
    else:
        # If no anchor, keep at top (insert after title)
        if len(doc.paragraphs) > 0:
            anchor_p = doc.paragraphs[0]._p
            anchor_idx = list(body).index(anchor_p)
            for i, el in enumerate(moved):
                body.insert(anchor_idx + 1 + i, el)

    doc.save(str(docx_path))


async def main():
    # Required CSVs must exist (per your instruction)
    if not RAW_CSV.exists():
        raise FileNotFoundError(f"Expected existing file not found: {RAW_CSV}")
    if not GROUPED_CSV.exists():
        raise FileNotFoundError(f"Expected existing file not found: {GROUPED_CSV}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    CACHE_DIR.mkdir(parents=True, exist_ok=True)

    rows = read_grouped_csv(GROUPED_CSV)

    # Validation cap (keep resume logic intact)
    if MAX_TO_SAVE is not None:
        rows = rows[:MAX_TO_SAVE]

    total = len(rows)

    state = load_state()
    last_completed = int(state.get("last_completed_design_seq", 0))
    cached_set = set(state.get("cached_design_seqs", []))

    print(f"Script dir:  {BASE_DIR}")
    print(f"Grouped CSV: {GROUPED_CSV.name}")
    print(f"Cache dir:   {CACHE_DIR}")
    print(f"State file:  {STATE_PATH}")
    print(f"DOCX out:    {DOCX_OUT.name}")
    print(f"Total now:   {total}")
    print(f"Resume: last_completed_design_seq={last_completed}, cached={len(cached_set)}\n")

    chrome_path = chrome_executable_path()
    print(f"Using local Chrome: {chrome_path}\n")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            executable_path=chrome_path,
            args=["--no-first-run", "--no-default-browser-check", "--disable-dev-shm-usage"],
        )
        context = await browser.new_context()

        for idx, row in enumerate(rows, start=1):
            # Resume rule: skip completed
            if row.design_seq <= last_completed:
                print(f"[{idx}/{total}] SKIP DONE  {row.design_seq:04d} :: {row.name}")
                continue

            cache_file = cache_text_path(row.design_seq, row.name)
            if cache_file.exists():
                print(f"[{idx}/{total}] CACHE HIT  {cache_file.name}")
                cached_set.add(row.design_seq)
            else:
                payload = await fetch_page_text(context, row, idx, total)
                written = write_cache(payload)
                cached_set.add(row.design_seq)
                print(f"[{idx}/{total}] CACHED     {written.name}")

            # checkpoint AFTER successful cache creation
            state["last_completed_design_seq"] = row.design_seq
            state["cached_design_seqs"] = sorted(cached_set)
            save_state(state)

            delay = random.randint(DELAY_MIN, DELAY_MAX)
            print(f"[{idx}/{total}] WAIT       {delay}s\n")
            await asyncio.sleep(delay)

        await context.close()
        await browser.close()

    # Build DOCX from cached content (idempotent rebuild)
    print("Building DOCX from cached pages...")
    cached_payloads = load_all_cached_payloads(rows)
    build_docx_from_cache(rows, cached_payloads, DOCX_TMP)

    # Insert TOC (created at end of process, then inserted into final document)
    print("Inserting TOC into DOCX...")
    DOCX_TMP.replace(DOCX_OUT)  # promote temp -> final
    insert_toc_at_front(DOCX_OUT)

    print("\nDone.")
    print(f"DOCX:  {DOCX_OUT}")
    print(f"State: {STATE_PATH}")
    print(f"Cache: {CACHE_DIR}")


if __name__ == "__main__":
    if sys.platform.startswith("win"):
        asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())
    asyncio.run(main())
